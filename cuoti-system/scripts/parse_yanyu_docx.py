#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
言语理解专项Word文档解析脚本

解析言语理解专项一.docx和解析文档，生成结构化JSON
"""
import os
import sys
import re
import json
from datetime import datetime

# 添加项目路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def parse_yanyu_questions(file_path):
    """解析言语理解题目文档"""
    from docx import Document
    
    doc = Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # 合并文本
    text_content = '\n'.join(full_text)
    
    # 打印前50行用于调试
    lines = text_content.split('\n')
    print(f"=== 文档共 {len(lines)} 行 ===")
    print("=== 前50行内容 ===")
    for i, line in enumerate(lines[:50]):
        print(f"{i+1}: {line}")
    
    return parse_questions_text(text_content)


def parse_questions_text(text):
    """解析题目文本"""
    questions = []
    lines = text.split('\n')
    
    # 题号匹配模式
    question_pattern = re.compile(r'^(\d+)[\.、．]\s*(.*)$')
    
    current_question = None
    stem_lines = []
    options_found = False
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        
        if not line:
            continue
        
        # 跳过标题行
        if '言语理解' in line or '第一部分' in line or '第二部分' in line or '片段阅读' in line or '逻辑填空' in line or '语句表达' in line:
            continue
        
        # 检查是否是新题目
        q_match = question_pattern.match(line)
        if q_match:
            # 保存上一题
            if current_question:
                # 将收集的stem_lines合并到题干
                if stem_lines and not current_question['stem']:
                    current_question['stem'] = ' '.join(stem_lines)
                elif stem_lines:
                    current_question['stem'] += ' ' + ' '.join(stem_lines)
                questions.append(current_question)
            
            # 开始新题目
            num = int(q_match.group(1))
            stem_start = q_match.group(2).strip()
            
            current_question = {
                'number': num,
                'stem': stem_start,
                'options': {'A': '', 'B': '', 'C': '', 'D': ''},
                'answer': '',
                'analysis': '',
                'category': '言语理解与表达',
                'subcategory': '',
                'knowledge_point': '',
                'difficulty': '中等'
            }
            stem_lines = []
            options_found = False
            continue
        
        if current_question:
            # 检查是否包含行内选项（A、xxx B、xxx 或 C、xxx D、xxx 格式）
            # 格式：A、内容                   B、内容（多个空格分隔）
            
            # 检测是否是选项行（包含 X、 格式）
            has_option_format = re.search(r'[A-D][、．\.]\s*.+', line)
            
            if has_option_format and not line.startswith('A项') and not line.startswith('B项') and not line.startswith('C项') and not line.startswith('D项'):
                options_found = True
                
                # 使用正则分割多个选项（按多个空格分隔）
                # 格式可能是: A、xxx     B、xxx 或 C、xxx     D、xxx
                parts = re.split(r'\s{3,}', line)  # 按3个或以上空格分割
                
                for part in parts:
                    part = part.strip()
                    if not part:
                        continue
                    
                    opt_match = re.match(r'^([A-D])[、．\.]\s*(.+)$', part)
                    if opt_match:
                        letter = opt_match.group(1)
                        content = opt_match.group(2).strip()
                        current_question['options'][letter] = content
            
            elif not options_found:
                # 还没遇到选项，继续收集题干
                stem_lines.append(line)
    
    # 保存最后一题
    if current_question:
        if stem_lines and not current_question['stem']:
            current_question['stem'] = ' '.join(stem_lines)
        elif stem_lines:
            current_question['stem'] += ' ' + ' '.join(stem_lines)
        questions.append(current_question)
    
    return questions




def parse_analysis_document(file_path):
    """解析解析文档，提取答案和解析"""
    from docx import Document
    
    doc = Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    text_content = '\n'.join(full_text)
    
    # 打印前30行用于调试
    lines = text_content.split('\n')
    print(f"\n=== 解析文档共 {len(lines)} 行 ===")
    print("=== 前30行内容 ===")
    for i, line in enumerate(lines[:30]):
        print(f"{i+1}: {line}")
    
    return parse_analysis_text(text_content)


def parse_analysis_text(text):
    """解析答案和解析文本
    
    注意：解析文档按模块分开编号
    - 片段阅读：1-40（对应题目1-40）
    - 语句表达：1-30（对应题目41-70）
    - 选词填空：1-40（对应题目71-110）
    """
    analysis_data = {}
    lines = text.split('\n')
    
    # 答案/解析行模式
    answer_pattern = re.compile(r'^(\d+)[、．\.]\s*正确答案是[：:]\s*([A-D])')
    
    # 模块偏移量
    current_offset = 0
    current_section = '片段阅读'
    
    current_number = None
    current_analysis = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 检测模块切换
        if '片段阅读' in line and '：' in line:
            current_section = '片段阅读'
            current_offset = 0
            continue
        elif '语句表达' in line and '：' in line:
            current_section = '语句表达'
            current_offset = 40  # 片段阅读40题后开始
            # 保存上一题的解析
            if current_number is not None and current_analysis:
                save_analysis(analysis_data, current_number, current_analysis)
            current_number = None
            current_analysis = []
            continue
        elif line.startswith('1、正确答案是') and current_section == '语句表达' and 30 in [n - 40 for n in analysis_data.keys() if n > 40]:
            # 检测到第三部分开始（选词填空）
            current_section = '选词填空'
            current_offset = 70  # 语句表达30题后开始
        
        # 跳过标题行
        if '言语' in line and '解析' in line and len(line) < 20:
            continue
        if line == '解析':
            continue
        
        # 检查是否是新的答案行
        match = answer_pattern.match(line)
        if match:
            # 保存上一题的解析
            if current_number is not None and current_analysis:
                save_analysis(analysis_data, current_number, current_analysis)
            
            section_num = int(match.group(1))
            answer = match.group(2)
            
            # 计算全局题号
            global_num = section_num + current_offset
            
            # 特殊处理：如果这个全局题号已经存在，说明进入了新模块
            if global_num in analysis_data and current_section != '选词填空':
                if current_section == '语句表达':
                    current_section = '选词填空'
                    current_offset = 70
                    global_num = section_num + current_offset
            
            current_number = global_num
            current_analysis = []
            
            if global_num not in analysis_data:
                analysis_data[global_num] = {'answer': '', 'analysis': ''}
            analysis_data[global_num]['answer'] = answer
        
        elif current_number is not None:
            # 继续收集解析内容
            current_analysis.append(line)
    
    # 保存最后一题的解析
    if current_number is not None and current_analysis:
        save_analysis(analysis_data, current_number, current_analysis)
    
    return analysis_data


def save_analysis(analysis_data, num, analysis_lines):
    """保存解析内容"""
    if num not in analysis_data:
        analysis_data[num] = {'answer': '', 'analysis': ''}
    
    # 过滤掉选项分析行，只保留核心解析
    filtered = []
    for a in analysis_lines:
        if not a.startswith('A项') and not a.startswith('B项') and not a.startswith('C项') and not a.startswith('D项'):
            filtered.append(a)
    
    analysis_data[num]['analysis'] = ' '.join(filtered) if filtered else ' '.join(analysis_lines)


def merge_questions_and_analysis(questions, analysis_data):
    """合并题目和解析数据"""
    matched = 0
    for q in questions:
        num = q['number']
        if num in analysis_data:
            if analysis_data[num]['answer']:
                q['answer'] = analysis_data[num]['answer']
                matched += 1
            if analysis_data[num]['analysis']:
                q['analysis'] = analysis_data[num]['analysis']
    
    print(f"成功匹配 {matched} 道题目的答案")
    
    # 打印解析数据的前5条用于验证
    print("\n=== 解析数据验证（前5条）===")
    for i in range(1, 6):
        if i in analysis_data:
            ans = analysis_data[i]['answer']
            ana = analysis_data[i]['analysis'][:50] if analysis_data[i]['analysis'] else '(无)'
            print(f"第{i}题: 答案={ans}, 解析={ana}...")
    
    return questions


def classify_question(question):
    """根据题目内容自动分类"""
    stem = question['stem']
    
    # 片段阅读类型判断（按提问方式分类）
    if '标题' in stem:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '标题添加题'
    elif '意在' in stem or '意图' in stem or '想要表达' in stem:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '意图判断题'
    elif '主旨' in stem or '主要说明' in stem or '重在强调' in stem or '核心观点' in stem:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '主旨概括题'
    elif '不符合' in stem or '不正确' in stem or '正确的是' in stem or '符合文意' in stem or '理解正确' in stem or '理解错误' in stem:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '细节理解题'
    elif '可以推出' in stem or '可以得知' in stem or '无法推断' in stem or '能推出' in stem:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '细节理解题'
    elif '这段文字' in stem or '这段话' in stem or '上述文字' in stem or '上述材料' in stem or '文段' in stem:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '片段阅读'
    
    # 逻辑填空（填空类）
    elif '____' in stem or '______' in stem or '（    ）' in stem or '（  ）' in stem:
        question['subcategory'] = '逻辑填空'
        # 判断是成语还是实词
        options_text = ' '.join(question['options'].values())
        if re.search(r'[\u4e00-\u9fff]{4}', options_text):  # 四字词语
            question['knowledge_point'] = '成语填空'
        else:
            question['knowledge_point'] = '实词填空'
    elif '填入' in stem and '最恰当' in stem:
        question['subcategory'] = '逻辑填空'
        question['knowledge_point'] = '逻辑填空'
    
    # 语句表达
    elif '排序' in stem or '顺序正确' in stem:
        question['subcategory'] = '语句表达'
        question['knowledge_point'] = '语句排序'
    elif '下文' in stem or '接下来' in stem:
        question['subcategory'] = '语句表达'
        question['knowledge_point'] = '下文推断'
    
    # 默认分类
    else:
        question['subcategory'] = '片段阅读'
        question['knowledge_point'] = '片段阅读'
    
    return question


def main():
    """主函数"""
    questions_file = '/Users/chaim/CodeBuddy/公考项目/docs/错题收集系统/言语理解专项一.docx'
    analysis_file = '/Users/chaim/CodeBuddy/公考项目/docs/错题收集系统/言语理解专项一解析.docx'
    output_file = '/Users/chaim/CodeBuddy/公考项目/cuoti-system/data/parsed/言语理解专项一/questions.json'
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    print("=" * 60)
    print("言语理解专项一 文档解析")
    print("=" * 60)
    
    # Step 1: 解析题目文档
    print("\n[Step 1] 解析题目文档...")
    questions = parse_yanyu_questions(questions_file)
    print(f"共解析到 {len(questions)} 道题目")
    
    # Step 2: 解析解析文档
    print("\n[Step 2] 解析解析文档...")
    analysis_data = parse_analysis_document(analysis_file)
    print(f"共解析到 {len(analysis_data)} 条解析")
    
    # Step 3: 合并数据
    print("\n[Step 3] 合并题目和解析...")
    questions = merge_questions_and_analysis(questions, analysis_data)
    
    # Step 4: 自动分类
    print("\n[Step 4] 自动分类...")
    for q in questions:
        classify_question(q)
    
    # Step 5: 统计信息
    with_answer = sum(1 for q in questions if q['answer'])
    with_analysis = sum(1 for q in questions if q['analysis'])
    print(f"\n统计信息:")
    print(f"  - 总题目数: {len(questions)}")
    print(f"  - 有答案: {with_answer}")
    print(f"  - 有解析: {with_analysis}")
    
    # 打印前3题作为示例
    print("\n=== 示例题目 ===")
    for i, q in enumerate(questions[:3]):
        print(f"\n--- 第 {q['number']} 题 ---")
        print(f"题干: {q['stem'][:100]}...")
        print(f"选项A: {q['options']['A'][:50] if q['options']['A'] else '(空)'}")
        print(f"答案: {q['answer'] or '(无)'}")
        print(f"分类: {q['subcategory'] or '(未分类)'}")
    
    # Step 6: 保存JSON
    output_data = {
        'source_file': os.path.basename(questions_file),
        'parse_time': datetime.now().isoformat(),
        'total_questions': len(questions),
        'category': '言语理解与表达',
        'source': '言语理解专项一',
        'questions': questions
    }
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n[完成] JSON已保存到: {output_file}")
    
    return output_data


if __name__ == '__main__':
    main()
